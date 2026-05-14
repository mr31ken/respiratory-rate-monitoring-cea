#!/usr/bin/env python3
"""
build_supplementary_docx.py — Generate submission-ready Supplementary DOCX.
Uses python-docx. Times New Roman 11pt, 1.5 spacing.
Reads supplementary.md and converts to DOCX programmatically.
"""

import os, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = Path(__file__).resolve().parent.parent
SUPP = ROOT / "supplementary"
OUTPUT = SUPP / "supplementary_20260405.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE = 11
LINE_SPACING = WD_LINE_SPACING.ONE_POINT_FIVE


def set_run_font(run, name=FONT_NAME, size=FONT_SIZE, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc, text="", bold=False, italic=False, alignment=None,
             space_after=None, space_before=None, first_line_indent=None,
             left_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    fmt = p.paragraph_format
    fmt.line_spacing_rule = LINE_SPACING
    if alignment:
        fmt.alignment = alignment
    if space_after is not None:
        fmt.space_after = Pt(space_after)
    if space_before is not None:
        fmt.space_before = Pt(space_before)
    if first_line_indent is not None:
        fmt.first_line_indent = Inches(first_line_indent)
    if left_indent is not None:
        fmt.left_indent = Inches(left_indent)
    return p


def add_mixed(doc, parts, alignment=None, first_line_indent=None,
              left_indent=None, space_before=None, space_after=None):
    """parts: list of (text, bold, italic)"""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_run_font(run, bold=bold, italic=italic)
    fmt = p.paragraph_format
    fmt.line_spacing_rule = LINE_SPACING
    if alignment:
        fmt.alignment = alignment
    if first_line_indent is not None:
        fmt.first_line_indent = Inches(first_line_indent)
    if left_indent is not None:
        fmt.left_indent = Inches(left_indent)
    if space_before is not None:
        fmt.space_before = Pt(space_before)
    if space_after is not None:
        fmt.space_after = Pt(space_after)
    return p


def add_heading_centered(doc, text):
    return add_para(doc, text, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_before=12, space_after=6)


def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=True, italic=True)
    fmt = p.paragraph_format
    fmt.line_spacing_rule = LINE_SPACING
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(3)
    return p


def add_body(doc, text, indent=True):
    return add_para(doc, text, first_line_indent=0.4 if indent else None)


def flush_table(doc, rows):
    if not rows:
        return
    headers = [c.strip() for c in rows[0].split("|") if c.strip()]
    data = []
    for row in rows[2:]:  # skip separator
        cells = [c.strip() for c in row.split("|") if c.strip()]
        if cells:
            data.append(cells)
    if not headers or not data:
        return
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(data), cols=ncols)
    table.style = "Table Grid"
    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=9)
    # Data rows
    for i, row_data in enumerate(data):
        for j in range(min(len(row_data), ncols)):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(row_data[j])
            set_run_font(run, size=9)


# ── Read markdown ────────────────────────────────────────────────────
md_text = (SUPP / "supplementary_20260405.md").read_text(encoding="utf-8")

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Page numbers (footer)
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
run.font.name = FONT_NAME
run.font.size = Pt(10)
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar1)
run2 = fp.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run2._r.append(instrText)
run3 = fp.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run3._r.append(fldChar2)

# Running header
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_h = hp.add_run("Supplementary Materials — Economic Evaluation of Automated RR Monitoring")
run_h.font.name = FONT_NAME
run_h.font.size = Pt(8)
run_h.italic = True

# Default style
style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.size = Pt(FONT_SIZE)
style.paragraph_format.line_spacing_rule = LINE_SPACING

# ══════════════════════════════════════════════════════════════════════
# Parse markdown into sections
# ══════════════════════════════════════════════════════════════════════
lines = md_text.split("\n")
sections_list = []
current_heading = ""
current_lines = []

for line in lines:
    if line.startswith("# ") and not line.startswith("## "):
        if current_heading or current_lines:
            sections_list.append((current_heading, current_lines))
        current_heading = line[2:].strip()
        current_lines = []
    elif line.startswith("## "):
        if current_heading or current_lines:
            sections_list.append((current_heading, current_lines))
        current_heading = line[3:].strip()
        current_lines = []
    else:
        current_lines.append(line)

if current_heading or current_lines:
    sections_list.append((current_heading, current_lines))

# ══════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════

add_para(doc, "SUPPLEMENTARY MATERIALS", bold=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para(doc, "", space_after=6)
add_para(doc, "Economic Evaluation of Automated Respiratory Rate Monitoring",
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_para(doc, "in General Hospital Wards:",
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_para(doc, "A Deterministic Cost-Minimization and Break-Even Modeling Study",
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# Process each section
# ══════════════════════════════════════════════════════════════════════

skip_sections = {
    "Supplementary Materials",
    "Economic Evaluation of Automated Respiratory Rate Monitoring in General Hospital Wards: A Deterministic Cost-Minimization and Break-Even Modeling Study",
}

for heading, content_lines in sections_list:
    if heading in skip_sections:
        continue
    if not heading and all(l.strip() in ("", "---") for l in content_lines):
        continue

    # Section heading
    if heading:
        add_heading_centered(doc, heading)

    # Process content
    in_table = False
    table_rows = []
    paragraph_buffer = ""

    def flush_paragraph(doc, text):
        if not text.strip():
            return
        text = text.strip()

        # Bold label pattern
        bold_match = re.match(r'\*\*(.+?)\*\*\s*(.*)', text)
        if bold_match:
            label = bold_match.group(1)
            rest = bold_match.group(2)
            if re.match(r'^\d+\.', label):
                add_mixed(doc, [(label, True, False), (" " + rest, False, False)],
                         left_indent=0.4)
            elif ":" in label:
                add_mixed(doc, [(label, True, False),
                               (" " + rest if rest else "", False, False)])
            else:
                add_mixed(doc, [(label, True, False),
                               (" " + rest if rest else "", False, False)])
        elif text.startswith("- **"):
            m = re.match(r'- \*\*(.+?)\*\*\s*(.*)', text)
            if m:
                add_mixed(doc, [(m.group(1), True, False),
                               (" " + m.group(2), False, False)],
                         left_indent=0.4)
            else:
                add_body(doc, text.lstrip("- "))
        elif text.startswith("- "):
            add_body(doc, text[2:], indent=False)
        elif text.startswith("[") and "]" in text:
            add_body(doc, text, indent=False)
        else:
            ref_match = re.match(r'^\[(\d+)\]\s+(.*)', text)
            if ref_match:
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_run_font(run)
                fmt = p.paragraph_format
                fmt.line_spacing_rule = LINE_SPACING
                fmt.left_indent = Inches(0.4)
                fmt.first_line_indent = Inches(-0.4)
            elif re.match(r'^\d+\.\s+\*\*', text):
                m = re.match(r'^(\d+\.\s+)\*\*(.+?)\*\*\s*(.*)', text)
                if m:
                    add_mixed(doc, [(m.group(1) + m.group(2), True, False),
                                   (" " + m.group(3), False, False)],
                             left_indent=0.4)
                else:
                    add_body(doc, text)
            else:
                add_body(doc, text, indent=True)

    for line in content_lines:
        stripped = line.strip()

        if stripped == "---":
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            continue

        if stripped.startswith("```"):
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            continue

        # Table detection
        if "|" in stripped and stripped.startswith("|"):
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(stripped)
            continue
        elif in_table:
            flush_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Subheading
        if stripped.startswith("### "):
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            add_subheading(doc, stripped[4:])
            continue

        if stripped.startswith("#### "):
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            add_subheading(doc, stripped[5:])
            continue

        # Empty line = paragraph break
        if not stripped:
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            continue

        # Accumulate
        if paragraph_buffer:
            paragraph_buffer += " " + stripped
        else:
            paragraph_buffer = stripped

    # Flush remaining
    if paragraph_buffer:
        flush_paragraph(doc, paragraph_buffer)
    if in_table and table_rows:
        flush_table(doc, table_rows)

    # Page break after each S-section
    if heading and heading.startswith("S"):
        doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════

doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
print(f"Size: {OUTPUT.stat().st_size / 1024:.0f} KB")
