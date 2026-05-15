#!/usr/bin/env python3
"""
build_docx.py — Generate submission-ready DOCX for Critical Care Medicine.
Uses python-docx. Times New Roman 12pt, double-spaced, line numbers, page numbers.
Figures inserted at end on individual pages.

Reads manuscript.md and converts to DOCX programmatically.
"""

import os, sys, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = Path(__file__).resolve().parent.parent
MANUSCRIPT = ROOT / "manuscript"
FIGURES = ROOT / "figures"
OUTPUT = MANUSCRIPT / "manuscript_20260516.docx"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc, text="", bold=False, italic=False, alignment=None,
             space_after=None, space_before=None, first_line_indent=None,
             left_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if alignment:
        fmt.alignment = alignment
    if space_after is not None:
        fmt.space_after = Pt(space_after)
    if space_before is not None:
        fmt.space_before = Pt(space_before)
    if first_line_indent is not None:
        fmt.first_line_indent = Inches(first_line_indent)
    if left_indent is not None:
        fmt.left_indent = Inches(left_indent)
    return p


def add_mixed(doc, parts, alignment=None, first_line_indent=None, left_indent=None,
              space_before=None, space_after=None):
    """parts: list of (text, bold, italic)"""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        set_run_font(run, bold=bold, italic=italic)
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if alignment:
        fmt.alignment = alignment
    if first_line_indent is not None:
        fmt.first_line_indent = Inches(first_line_indent)
    if left_indent is not None:
        fmt.left_indent = Inches(left_indent)
    if space_before is not None:
        fmt.space_before = Pt(space_before)
    if space_after is not None:
        fmt.space_after = Pt(space_after)
    return p


def add_heading_centered(doc, text):
    p = add_para(doc, text, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=12, space_after=6)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=True, italic=True)
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(3)
    return p


def add_body(doc, text, indent=True):
    return add_para(doc, text, first_line_indent=0.5 if indent else None)


def add_page_break(doc):
    doc.add_page_break()


# ── Parse markdown and build document ────────────────────────────────

md_text = (MANUSCRIPT / "manuscript_20260516.md").read_text(encoding="utf-8")

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Line numbers
sectPr = section._sectPr
lnNumType = parse_xml(f'<w:lnNumType {nsdecls("w")} w:countBy="1" w:restart="continuous"/>')
sectPr.append(lnNumType)

# Page numbers (footer)
footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
run.font.name = "Times New Roman"
run.font.size = Pt(10)
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar1)
run2 = fp.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run2._r.append(instrText)
run3 = fp.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run3._r.append(fldChar2)

# Running header
header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_h = hp.add_run("Economic Evaluation of Automated RR Monitoring")
run_h.font.name = "Times New Roman"
run_h.font.size = Pt(9)
run_h.italic = True

# Default style
style = doc.styles['Normal']
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# ══════════════════════════════════════════════════════════════════════
# Parse markdown sections
# ══════════════════════════════════════════════════════════════════════

lines = md_text.split("\n")

# State machine: track current section
current_section = None
in_table = False
table_rows = []
skip_metadata = True  # Skip lines until after first ---

# Helper to flush table
def flush_table(doc, rows):
    if not rows:
        return
    # Parse header and data
    headers = [c.strip() for c in rows[0].split("|") if c.strip()]
    data = []
    for row in rows[2:]:  # skip separator
        cells = [c.strip() for c in row.split("|") if c.strip()]
        if cells:
            data.append(cells)
    if not headers or not data:
        return

    ncols = len(headers)
    table = doc.add_table(rows=1 + len(data), cols=ncols)
    table.style = "Table Grid"

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=9)

    # Data rows
    for i, row_data in enumerate(data):
        for j in range(min(len(row_data), ncols)):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(row_data[j])
            set_run_font(run, size=9)


i = 0
section_count = 0

# First pass: find section boundaries
sections = []
current_lines = []
current_heading = ""

for line in lines:
    if line.startswith("# ") and not line.startswith("## "):
        # Title
        if current_heading or current_lines:
            sections.append((current_heading, current_lines))
        current_heading = line[2:].strip()
        current_lines = []
    elif line.startswith("## "):
        if current_heading or current_lines:
            sections.append((current_heading, current_lines))
        current_heading = line[3:].strip()
        current_lines = []
    else:
        current_lines.append(line)

if current_heading or current_lines:
    sections.append((current_heading, current_lines))

# ══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════

add_para(doc, "", space_after=72)
add_para(doc, "Economic Evaluation of Automated Respiratory Rate Monitoring",
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_para(doc, "in General Hospital Wards:",
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
add_para(doc, "A Deterministic Cost-Minimization and Break-Even Modeling Study",
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

# Pull Authors section content from parsed sections (single source of truth: manuscript_20260516.md)
authors_section = next(((h, c) for h, c in sections if h == "Authors"), None)
if authors_section is not None:
    author_lines = [l.strip() for l in authors_section[1] if l.strip()]
    # First non-empty line: author names with superscripts
    # Subsequent lines starting with superscript char: affiliations
    # Last line(s): "Corresponding author: ..." block (starts with **)
    for line in author_lines:
        if line.startswith("**Corresponding author:**"):
            # Render bold label + rest, full-width, centered
            rest = line.replace("**Corresponding author:**", "").strip()
            # Strip markdown link syntax [text](url) → text
            rest = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', rest)
            add_mixed(doc, [("Corresponding author: ", True, False),
                            (rest, False, False)],
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=24)
        elif line and line[0] in "¹²³⁴⁵⁶⁷⁸⁹":
            # Affiliation line — centered, no indent
            add_para(doc, line, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        else:
            # Author name line — centered, slightly larger feel via spacing
            add_para(doc, line, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
else:
    add_para(doc, "[Authors section missing in manuscript_20260516.md]",
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_para(doc, "Tables: 3  |  Figures: 6  |  References: 25",
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=0)
add_para(doc, "Supplementary Materials: 10 sections",
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════
# Process each section from parsed markdown
# ══════════════════════════════════════════════════════════════════════

skip_sections = {"Authors", "Target Journal"}

for heading, content_lines in sections:
    if heading in skip_sections:
        continue

    # Skip separator-only sections
    if not heading and all(l.strip() in ("", "---") for l in content_lines):
        continue

    # Section heading
    if heading:
        heading_upper = heading.upper()
        if heading_upper in ("ABSTRACT", "INTRODUCTION", "METHODS", "RESULTS",
                            "DISCUSSION", "CONCLUSIONS", "DECLARATIONS",
                            "REFERENCES", "FIGURE LEGENDS", "TABLES"):
            add_heading_centered(doc, heading_upper)
        elif heading.startswith("###"):
            # Sub-subheading
            add_subheading(doc, heading.replace("### ", ""))
        else:
            add_heading_centered(doc, heading)

    # Process content lines
    in_table = False
    table_rows = []
    paragraph_buffer = ""

    def flush_paragraph(doc, text):
        if not text.strip():
            return
        text = text.strip()

        # Handle **bold label:** pattern at start
        bold_match = re.match(r'\*\*(.+?)\*\*\s*(.*)', text)
        if bold_match:
            label = bold_match.group(1)
            rest = bold_match.group(2)
            # Check if it's a list item (starts with number or bullet)
            if re.match(r'^\d+\.', label):
                add_mixed(doc, [(label, True, False), (" " + rest, False, False)],
                         left_indent=0.5)
            elif ":" in label:
                add_mixed(doc, [(label, True, False), (" " + rest if rest else "", False, False)],
                         first_line_indent=0.5 if not rest else None)
            else:
                add_mixed(doc, [(label, True, False), (" " + rest if rest else "", False, False)])
        elif text.startswith("- **"):
            # Bullet with bold label
            m = re.match(r'- \*\*(.+?)\*\*\s*(.*)', text)
            if m:
                add_mixed(doc, [(m.group(1), True, False), (" " + m.group(2), False, False)],
                         left_indent=0.5)
            else:
                add_body(doc, text.lstrip("- "))
        elif text.startswith("- "):
            add_body(doc, text[2:], indent=False)
        elif text.startswith("[") and text.endswith(")"):
            # Reference line - skip markdown links
            add_body(doc, text, indent=False)
        else:
            # Check if starts with a number (numbered list or reference)
            ref_match = re.match(r'^\[(\d+)\]\s+(.*)', text)
            if ref_match:
                # Reference
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_run_font(run)
                fmt = p.paragraph_format
                fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                fmt.left_indent = Inches(0.5)
                fmt.first_line_indent = Inches(-0.5)
            elif re.match(r'^\d+\.\s+\*\*', text):
                # Numbered list with bold
                m = re.match(r'^(\d+\.\s+)\*\*(.+?)\*\*\s*(.*)', text)
                if m:
                    add_mixed(doc, [(m.group(1) + m.group(2), True, False),
                                   (" " + m.group(3), False, False)],
                             left_indent=0.5)
                else:
                    add_body(doc, text)
            else:
                add_body(doc, text, indent=True)

    for line in content_lines:
        stripped = line.strip()

        if stripped == "---":
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            continue

        if stripped.startswith("```"):
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            continue

        # Table detection
        if "|" in stripped and stripped.startswith("|"):
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(stripped)
            continue
        elif in_table:
            flush_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Subheading
        if stripped.startswith("### "):
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            add_subheading(doc, stripped[4:])
            continue

        if stripped.startswith("#### "):
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            add_subheading(doc, stripped[5:])
            continue

        # Empty line = paragraph break
        if not stripped:
            if paragraph_buffer:
                flush_paragraph(doc, paragraph_buffer)
                paragraph_buffer = ""
            continue

        # Accumulate paragraph
        if paragraph_buffer:
            paragraph_buffer += " " + stripped
        else:
            paragraph_buffer = stripped

    # Flush remaining
    if paragraph_buffer:
        flush_paragraph(doc, paragraph_buffer)
    if in_table and table_rows:
        flush_table(doc, table_rows)

    # Page break after major sections
    if heading and heading.upper() in ("ABSTRACT", "INTRODUCTION", "METHODS",
                                        "RESULTS", "DISCUSSION", "CONCLUSIONS",
                                        "DECLARATIONS", "REFERENCES",
                                        "FIGURE LEGENDS", "TABLES"):
        add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════
# FIGURES (each on its own page)
# ══════════════════════════════════════════════════════════════════════

figure_files = [
    ("fig0_study_overview.png", "Figure 1"),
    ("fig1_accuracy_comparison.png", "Figure 2"),
    ("fig2_cost_comparison.png", "Figure 3"),
    ("fig3_breakeven.png", "Figure 4"),
    ("fig4_tornado.png", "Figure 5"),
    ("fig5_scenario_analysis.png", "Figure 6"),
]

for fname, label in figure_files:
    fpath = FIGURES / fname
    if fpath.exists():
        add_heading_centered(doc, label)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(fpath), width=Inches(6.0))
        add_page_break(doc)
    else:
        add_para(doc, f"[{label}: {fname} not found]", italic=True)
        add_page_break(doc)

# ══════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════

doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
print(f"Size: {OUTPUT.stat().st_size / 1024:.0f} KB")
